import streamlit as st
from docx import Document

def generate_index(thesis_title):
    chapters = {
        "1. Introduction": ["1.1 Background", "1.2 Research Objectives", "1.3 Scope", "1.4 Methodology"],
        "2. Literature Review": ["2.1 Previous Research", "2.2 Theoretical Framework", "2.3 Research Gaps"],
        "3. Research Methodology": ["3.1 Data Collection", "3.2 Analysis Methods", "3.3 Limitations"],
        "4. Data Analysis": ["4.1 Data Presentation", "4.2 Interpretation of Results"],
        "5. Findings and Discussion": ["5.1 Key Findings", "5.2 Implications", "5.3 Comparison with Literature"],
        "6. Conclusion and Recommendations": ["6.1 Summary of Findings", "6.2 Recommendations", "6.3 Future Research Directions"],
        "7. References and Appendices": ["7.1 References", "7.2 Appendices"]
    }
    return chapters

def save_as_word(thesis_title, chapters):
    doc = Document()
    doc.add_heading(f"{thesis_title} - Index Page", level=1)
    
    for chapter, subchapters in chapters.items():
        doc.add_heading(chapter, level=2)
        for sub in subchapters:
            doc.add_paragraph(sub, style='List Bullet')
    
    filename = "Thesis_Index.docx"
    doc.save(filename)
    return filename

# Streamlit App
st.title("Thesis Index Generator")
st.write("Enter your thesis title to generate an index page with chapters and subchapters.")

# Thesis Title Input
thesis_title = st.text_input("Enter Thesis Title:")

if thesis_title:
    chapters = generate_index(thesis_title)
    
    st.write("## Generated Index Page")
    st.write(f"### {thesis_title}")
    
    for chapter, subchapters in chapters.items():
        st.subheader(chapter)
        for sub in subchapters:
            st.write(f"- {sub}")
    
    # Save as Word
    if st.button("Download as Word File"):
        filename = save_as_word(thesis_title, chapters)
        with open(filename, "rb") as file:
            st.download_button("Click to Download", file, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
